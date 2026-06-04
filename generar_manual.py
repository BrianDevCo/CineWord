#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ROJO = RGBColor(0xCC, 0x12, 0x44)
NEGRO = RGBColor(0x1A, 0x1A, 0x1A)
GRIS = RGBColor(0x55, 0x55, 0x55)
GRIS_CLARO = RGBColor(0xAA, 0xAA, 0xAA)
VERDE = RGBColor(0x16, 0xA3, 0x4A)
AZUL = RGBColor(0x1D, 0x4E, 0xD8)

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

# ── Estilos base ──────────────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = NEGRO

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = ROJO
    run.font.name = 'Calibri'
    # línea inferior
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'CC1244')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NEGRO
    run.font.name = 'Calibri'
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ROJO
    run.font.name = 'Calibri'
    return p

def body(text, bold=False, color=None, size=11, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color or NEGRO
    run.font.name = 'Calibri'
    return p

def paso(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"Paso {num}. ")
    r1.bold = True
    r1.font.color.rgb = ROJO
    r1.font.size = Pt(12)
    r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.bold = True
    r2.font.color.rgb = NEGRO
    r2.font.size = Pt(12)
    r2.font.name = 'Calibri'
    return p

def bullet(text, level=1, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(level * 0.7)
    p.paragraph_format.space_after  = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True
        r1.font.color.rgb = NEGRO
        r1.font.size = Pt(11)
        r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = NEGRO
    r2.font.name = 'Calibri'
    return p

def nota(text, tipo="info"):
    color_map = {"info": (0xDB, 0xEA, 0xFE), "warn": (0xFF, 0xED, 0xC2), "ok": (0xD1, 0xFA, 0xE5)}
    text_color_map = {"info": AZUL, "warn": RGBColor(0x92, 0x40, 0x00), "ok": VERDE}
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    icon = {"info": "ℹ️  ", "warn": "⚠️  ", "ok": "✅  "}[tipo]
    run = p.add_run(icon + text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = text_color_map[tipo]
    run.font.name = 'Calibri'
    run.italic = True
    return p

def separador():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("─" * 70)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    return p

def tabla_simple(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # cabecera
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'CC1244')
        tcPr.append(shd)
    # filas
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
#   PORTADA
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CINEWORLD")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = ROJO
r.font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Manual de Administración y Configuración")
r2.font.size = Pt(18)
r2.font.color.rgb = NEGRO
r2.font.name = 'Calibri'
r2.bold = True

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Centro Comercial MR Outlet · Cali, Colombia")
r3.font.size = Pt(12)
r3.font.color.rgb = GRIS
r3.font.name = 'Calibri'

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("Versión 1.0  —  Junio 2026")
r4.font.size = Pt(11)
r4.font.color.rgb = GRIS_CLARO
r4.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

# índice rápido
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run("CONTENIDO DE ESTE DOCUMENTO")
r5.bold = True
r5.font.size = Pt(11)
r5.font.color.rgb = GRIS
r5.font.name = 'Calibri'

contenido = [
    "1. Cómo usar la página web (guía para todos)",
    "2. Panel de Administración — guía detallada",
    "3. Configurar cuenta MercadoPago (Contabilidad)",
    "4. Configurar cuenta Resend — Emails automáticos",
    "5. Preguntas frecuentes y solución de problemas",
]
for item in contenido:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(item)
    r.font.size = Pt(11)
    r.font.color.rgb = NEGRO
    r.font.name = 'Calibri'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#   SECCIÓN 1 — CÓMO USAR LA PÁGINA WEB
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Cómo usar la página web")
body("Esta sección explica todo lo que puede hacer un visitante o cliente en el sitio web del cine. Está escrita para que cualquier persona del equipo pueda entenderla sin necesidad de conocimientos técnicos.")

h2("1.1  Inicio (Cartelera)")
body("Al entrar al sitio, los clientes ven directamente la cartelera con todas las películas disponibles. Cada película muestra:")
bullet("El póster y nombre de la película")
bullet("El género y duración")
bullet("Un botón para comprar boletas")
nota("Si una película no aparece, es porque fue desactivada desde el panel de administración.", "warn")

h2("1.2  Comprar boletas — Paso a paso del cliente")

paso(1, "Seleccionar la película")
body("El cliente hace clic en 'Comprar boletas' en la película que le interesa.", indent=True)

paso(2, "Elegir fecha y formato")
body("Aparece un calendario con los días disponibles. El cliente elige el día y el formato (2D, 3D, etc.). Solo aparecen fechas con funciones activas.", indent=True)

paso(3, "Elegir la hora")
body("Se muestran las funciones disponibles para ese día. Las funciones que ya pasaron no aparecen. El cliente selecciona la hora que prefiere.", indent=True)

paso(4, "Seleccionar las sillas")
body("Se abre un mapa interactivo de la sala. Las sillas en rojo ya están ocupadas. Las verdes están disponibles. El cliente hace clic para elegir su(s) silla(s).", indent=True)
nota("Las sillas seleccionadas quedan reservadas por 5 minutos mientras el cliente paga. Si no paga a tiempo, se liberan automáticamente.", "warn")

paso(5, "Ingresar datos y pagar")
body("El cliente ingresa su nombre, correo y teléfono. Luego paga con tarjeta débito, crédito o PSE a través de MercadoPago.", indent=True)

paso(6, "Confirmación y boleta por correo")
body("Después del pago exitoso, el sistema envía automáticamente un correo con la boleta y un código QR para entrar al cine.", indent=True)
nota("El correo llega en menos de 2 minutos. Si no aparece, revisar la carpeta de spam.", "info")

h2("1.3  Crear cuenta / Iniciar sesión")
body("Los clientes pueden crear una cuenta para guardar su historial de compras. Pueden registrarse con:")
bullet("Correo electrónico y contraseña")
bullet("Cuenta de Google (un clic)")
body("Con la cuenta activa, en la sección 'Mi cuenta' pueden ver todas sus reservas anteriores y descargar boletas.")

h2("1.4  Snacks y combos")
body("La sección de Snacks muestra el menú de confitería del cine (palomitas, bebidas, combos). Por ahora es informativa — los snacks se compran directamente en taquilla, no a través del sitio.")

h2("1.5  Sección de Contacto")
body("El sitio tiene una sección de contacto con la dirección, teléfono y horarios del cine. También hay un formulario de PQR (Peticiones, Quejas y Reclamos) en la página /pqr.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#   SECCIÓN 2 — PANEL DE ADMINISTRACIÓN
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Panel de Administración")
body("El panel de administración es la herramienta principal para gestionar el cine desde el sitio web. Solo puede acceder quien tenga los permisos de administrador.")

h2("2.1  Cómo acceder al panel")
bullet("Abrir el navegador e ir a:")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.4)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("https://cine-word.vercel.app/admin")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = AZUL
r.font.name = 'Calibri'
bullet("Iniciar sesión con el correo de administrador (brianl280499@gmail.com o el que asigne el desarrollador).")
bullet("Si no ha iniciado sesión, el sitio lo redirige automáticamente a la página de login.")
nota("Solo los correos autorizados pueden ver el panel. Un cliente normal que intente entrar a /admin será bloqueado.", "ok")

h2("2.2  Dashboard (Pantalla principal)")
body("Al entrar al panel, aparece el Dashboard con un resumen del negocio en tiempo real:")

tabla_simple(
    ["Indicador", "Qué significa"],
    [
        ["Películas en cartelera", "Cuántas películas están activas y disponibles para compra"],
        ["Próximos estrenos", "Películas cargadas con estado 'Próximo' (no disponibles aún para compra)"],
        ["Funciones próximas 7d", "Cuántas funciones están programadas en los próximos 7 días"],
        ["Reservas hoy", "Boletas vendidas el día de hoy"],
        ["Ingresos hoy", "Dinero recaudado hoy por ventas en línea"],
        ["Ingresos del mes", "Dinero total recaudado en el mes actual"],
    ],
    [5, 9.5]
)

body("También hay botones de acceso rápido para las acciones más comunes: Nueva película, Agregar función, Ver reservas y Gestionar salas.")

separador()

h2("2.3  Módulo: Películas")
body("Ruta: /admin/peliculas")
body("Aquí se administran todas las películas del cine — tanto las que están en cartelera como los próximos estrenos.")

h3("¿Qué se puede hacer?")
bullet("Ver la lista de todas las películas con su estado (En cartelera / Próximo / Retirada)")
bullet("Activar o desactivar una película (si se desactiva, desaparece del sitio)")
bullet("Crear una nueva película manualmente")
bullet("Editar los datos de una película existente")
bullet("Enriquecer con TMDB: busca automáticamente el póster, sinopsis y datos desde la base de datos internacional de películas")

h3("Cómo agregar una nueva película manualmente")
paso(1, "Clic en el botón rojo 'NUEVA PELÍCULA'")
paso(2, "Llenar el formulario:")
tabla_simple(
    ["Campo", "Descripción", "Obligatorio"],
    [
        ["Título", "Nombre exacto de la película", "Sí"],
        ["Estado", "En cartelera / Próximo / Retirada", "Sí"],
        ["Género", "Acción, Drama, Comedia, etc.", "No"],
        ["Duración", "En minutos (ej: 120)", "No"],
        ["Sinopsis", "Descripción corta de la película", "No"],
        ["Imagen (poster)", "URL de la imagen del póster", "No"],
        ["Clasificación", "G, PG, PG-13, R, etc.", "No"],
    ],
    [3.5, 7, 2.5]
)
paso(3, "Clic en 'GUARDAR' — la película aparece de inmediato en el listado")

h3("Botón 'ENRIQUECER CON TMDB'")
body("Este botón busca automáticamente en la base de datos de películas del mundo (TMDB) los pósters, sinopsis y datos de todas las películas que no tengan imagen. Es muy útil para completar info rápidamente sin tener que buscar imágenes manualmente.")
nota("Si la película no se encuentra automáticamente (por nombre extraño o película local), se puede ingresar el ID de TMDB manualmente en el campo que aparece junto a cada película.", "info")

separador()

h2("2.4  Módulo: Funciones")
body("Ruta: /admin/funciones")
body("Las funciones son las programaciones de cada película: en qué sala, qué día y a qué hora. Aquí se programan y gestionan todas las funciones.")

h3("¿Qué se puede hacer?")
bullet("Ver todas las funciones organizadas por fecha")
bullet("Crear una nueva función manualmente")
bullet("Activar o desactivar una función (toggle rojo/gris)")
bullet("Eliminar una función")
bullet("Sincronizar con Score (importar funciones del sistema del cine)")

h3("Cómo crear una nueva función")
paso(1, "Clic en 'NUEVA FUNCIÓN' (botón rojo)")
paso(2, "Llenar los campos:")
tabla_simple(
    ["Campo", "Qué poner"],
    [
        ["Película", "Seleccionar de la lista (solo aparecen películas activas)"],
        ["Sala", "La sala donde se proyectará"],
        ["Fecha", "El día de la función"],
        ["Hora", "La hora exacta de inicio"],
        ["Formato", "2D, 3D o VIP"],
        ["Precio regular", "Valor en pesos colombianos (ej: 14000)"],
        ["Precio VIP", "Precio para sillas VIP si aplica"],
    ],
    [3.5, 10]
)
paso(3, "Clic en 'CREAR FUNCIÓN'")

h3("Sincronizar con Score")
body("Score es el sistema POS físico del cine. Al hacer clic en 'SINCRONIZAR SCORE', el sistema importa automáticamente las películas y funciones desde Score, actualizando la cartelera del sitio web.")
nota("Esta sincronización debe hacerse cada jueves cuando llegue la cartelera nueva de la semana. Tarda menos de 30 segundos.", "info")
nota("Para que funcione, el computador del cine debe tener el programa ngrok corriendo. Si la sincronización falla, verificar que ngrok esté activo.", "warn")

separador()

h2("2.5  Módulo: Reservas")
body("Ruta: /admin/reservas")
body("Aquí se pueden ver todas las compras realizadas a través del sitio web.")

h3("¿Qué se puede ver?")
bullet("Lista completa de reservas con nombre del cliente, película, fecha, asientos y total pagado")
bullet("Estado de cada reserva: Pagado, Pendiente o Cancelado")
bullet("Total de ingresos del período filtrado")

h3("Cómo filtrar reservas")
bullet("Por estado: Todos / Pagado / Pendiente / Cancelado")
bullet("Por fecha: seleccionar rango 'Desde' y 'Hasta'")
bullet("Los resultados se actualizan automáticamente al cambiar los filtros")

tabla_simple(
    ["Estado", "Significado"],
    [
        ["Pagado ✅", "El cliente completó el pago — la boleta fue enviada por correo"],
        ["Pendiente ⏳", "El pago está en proceso o no se confirmó (MercadoPago aún procesando)"],
        ["Cancelado ❌", "El pago falló o fue cancelado por el cliente"],
    ],
    [4, 10]
)
nota("Solo las reservas con estado 'Pagado' representan ingresos reales.", "warn")

separador()

h2("2.6  Módulo: Salas")
body("Ruta: /admin/salas")
body("Aquí se configura la información de cada sala del cine.")
bullet("Nombre de la sala (Sala 2, Sala 4, etc.)")
bullet("Tipo (estándar, VIP, etc.)")
bullet("Capacidad y distribución de sillas")
nota("Normalmente las salas no cambian. Este módulo se usa principalmente cuando se incorpora una sala nueva o se modifica la distribución de asientos.", "info")

separador()

h2("2.7  Módulo: Hero (Imagen principal del sitio)")
body("Ruta: /admin/hero")
body("Permite cambiar las imágenes que aparecen en el banner principal (Hero) de la página de inicio. Aquí se pueden actualizar las imágenes promocionales que el cliente ve al entrar al sitio.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#   SECCIÓN 3 — MERCADOPAGO
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Configurar cuenta MercadoPago")
body("MercadoPago es la plataforma de pagos que usa el sitio web para procesar las compras de boletas. Contabilidad debe crear la cuenta oficial del cine para que los pagos lleguen directamente al negocio.")
nota("Actualmente el sitio está en modo de prueba. Hasta que no se configuren las credenciales de producción, los pagos NO son reales.", "warn")

h2("3.1  Crear la cuenta del negocio")
paso(1, "Abrir el navegador y entrar a: mercadopago.com.co")
paso(2, "Hacer clic en 'Crear cuenta'")
paso(3, "Seleccionar 'Soy una empresa' (no cuenta personal)")
paso(4, "Ingresar los datos del negocio:")
bullet("NIT de la empresa (o cédula del representante legal)")
bullet("Correo electrónico corporativo del cine")
bullet("Contraseña segura (mínimo 8 caracteres, con mayúsculas y números)")
paso(5, "Revisar el correo registrado y hacer clic en el enlace de verificación que llegó")

h2("3.2  Verificar la identidad (obligatorio para recibir pagos)")
paso(1, "Iniciar sesión en MercadoPago")
paso(2, "Ir a 'Mi perfil' → 'Activar mi cuenta'")
paso(3, "Subir los documentos solicitados:")
bullet("Cédula del representante legal (foto de ambos lados, legible)")
bullet("RUT de la empresa")
bullet("Cámara de comercio (puede ser requerida)")
paso(4, "Esperar la aprobación de MercadoPago")
nota("La verificación tarda entre 24 y 72 horas hábiles. MercadoPago envía un correo cuando está lista.", "info")

h2("3.3  Configurar cuenta bancaria para recibir retiros")
paso(1, "Ir a 'Mi dinero' → 'Cuentas bancarias'")
paso(2, "Clic en 'Agregar cuenta bancaria'")
paso(3, "Ingresar los datos bancarios:")
tabla_simple(
    ["Dato", "Descripción"],
    [
        ["Banco", "Nombre del banco (Bancolombia, Davivienda, etc.)"],
        ["Tipo de cuenta", "Ahorros o Corriente"],
        ["Número de cuenta", "El número exacto de la cuenta"],
        ["Titular", "Debe coincidir con el NIT o cédula registrada en MercadoPago"],
    ],
    [4, 10]
)
paso(4, "MercadoPago realizará un depósito de verificación de $1 COP a la cuenta. Confirmar el monto exacto recibido en la app.")
paso(5, "La cuenta bancaria quedará vinculada y lista para recibir retiros.")
nota("El titular de la cuenta bancaria DEBE ser el mismo que aparece en MercadoPago. Si son diferentes, el retiro será rechazado.", "warn")

h2("3.4  Activar retiros automáticos (recomendado)")
paso(1, "Ir a 'Mi dinero' → 'Retiros'")
paso(2, "Activar 'Retiro automático'")
paso(3, "Seleccionar la frecuencia: diario, semanal o mensual")
body("Con esta opción activa, el dinero de las ventas se transfiere automáticamente a la cuenta bancaria sin necesidad de hacerlo manualmente cada vez.")

h2("3.5  Obtener las credenciales para el desarrollador")
body("Este es el paso más importante para conectar la cuenta del cine con el sitio web. Sin estas credenciales, los pagos no llegan a la cuenta correcta.")
nota("IMPORTANTE: estas claves son confidenciales. Enviarlas únicamente al desarrollador por un canal seguro (correo o mensaje directo). NO compartir por grupos de WhatsApp.", "warn")
paso(1, "Ir a: mercadopago.com.co/developers/panel")
paso(2, "Hacer clic en 'Crear aplicación'")
paso(3, "En el formulario, ingresar:")
bullet("Nombre de la aplicación: CineWorld Web")
bullet("Tipo: Checkout Pro / Payments")
bullet("Clic en 'Guardar'")
paso(4, "Dentro de la aplicación creada, ir a 'Credenciales de producción'")
paso(5, "Copiar y enviar al desarrollador los dos valores:")
tabla_simple(
    ["Credencial", "Cómo luce", "Para qué sirve"],
    [
        ["Public Key", "APP_USR-abc123...", "Identifica la cuenta en el sitio web"],
        ["Access Token", "APP_USR-xyz789...", "Autoriza cobros reales a la cuenta"],
    ],
    [3.5, 5, 5.5]
)
nota("Estas credenciales solo funcionan para pagos reales (producción). Las actuales del sitio son de prueba.", "info")

h2("3.6  Configurar métodos de pago aceptados")
paso(1, "Dentro de la aplicación en el panel de desarrolladores, ir a 'Configuración'")
paso(2, "Activar los métodos de pago deseados:")
bullet("Tarjetas de crédito (Visa, Mastercard, Amex)")
bullet("Tarjetas débito")
bullet("PSE (transferencia bancaria)")
bullet("Efecty / Baloto (pago en efectivo — opcional)")
paso(3, "Guardar los cambios")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#   SECCIÓN 4 — RESEND
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Configurar Resend — Emails automáticos")
body("Resend es el servicio que usa el sitio web para enviar automáticamente las boletas por correo electrónico a los clientes después de cada compra. Esta configuración la hace el desarrollador, pero contabilidad o el administrador deben proveer el dominio del correo del cine.")
nota("Actualmente el sistema solo puede enviar correos a brianl280499@gmail.com (modo de prueba). Para enviar a cualquier cliente, se necesita configurar el dominio propio del cine.", "warn")

h2("4.1  Qué hace Resend en el sitio")
bullet("Después de cada pago exitoso, envía automáticamente un correo al cliente")
bullet("El correo contiene: resumen de la compra, nombre de la película, fecha, hora, sala, asientos y un código QR")
bullet("El cliente presenta ese QR en la puerta del cine para ingresar")
bullet("No hay necesidad de hacer nada manualmente — todo es automático")

h2("4.2  Lo que se necesita para activarlo en producción")
body("Para que los correos lleguen a TODOS los clientes (no solo al desarrollador), se necesita un dominio de correo del cine. Por ejemplo: boletas@cineworld.com.co")

tabla_simple(
    ["Requisito", "Quién lo consigue", "Descripción"],
    [
        ["Dominio propio", "Administración / Contabilidad", "Un dominio web del cine (ej: cineworld.com.co) comprado en cualquier registrador"],
        ["Acceso al DNS del dominio", "Quien administra el dominio", "Para agregar unos registros técnicos que verifican el dominio en Resend"],
        ["Correo de envío definido", "Administración", "El correo desde el que se enviarán las boletas (ej: boletas@cineworld.com.co)"],
    ],
    [4, 4, 6]
)

h2("4.3  Pasos para configurar Resend (los hace el desarrollador con info del cine)")
paso(1, "El desarrollador crea una cuenta en resend.com con el correo de desarrollo")
paso(2, "Va a 'Domains' → 'Add Domain' e ingresa el dominio del cine")
paso(3, "Resend entrega unos registros DNS (TXT, MX) que deben agregarse al dominio")
paso(4, "El administrador del dominio agrega esos registros (proceso de 5 minutos en el panel del registrador)")
paso(5, "Resend verifica el dominio automáticamente (tarda entre 15 minutos y 24 horas)")
paso(6, "El desarrollador actualiza la variable de entorno en Vercel con la nueva dirección de correo")
paso(7, "Listo — a partir de ese momento todos los clientes reciben sus boletas por correo")

h2("4.4  ¿Qué pasa si un cliente no recibe el correo?")
bullet("Pedir al cliente que revise su carpeta de spam o correo no deseado")
bullet("Verificar en el panel de administración → Reservas que el estado sea 'Pagado'")
bullet("Si el pago está confirmado pero el correo no llegó, el desarrollador puede reenviar la boleta manualmente desde Resend o Supabase")
nota("El correo QR es solo un respaldo. Si el cliente no lo recibe, puede mostrar su comprobante de pago de MercadoPago en taquilla.", "info")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#   SECCIÓN 5 — FAQ
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Preguntas frecuentes y solución de problemas")

h2("5.1  Preguntas del sitio web")

h3("¿Por qué no aparece una película en el sitio?")
body("Puede estar desactivada. Ir al panel Admin → Películas y verificar que el toggle esté en verde (activo).")

h3("¿Por qué no hay funciones disponibles para una película?")
body("Puede ser que:")
bullet("No se han creado funciones para esa película en Admin → Funciones")
bullet("Las funciones creadas ya pasaron (solo aparecen horas futuras)")
bullet("Las funciones están desactivadas (toggle en gris)")

h3("¿Por qué aparece 'No hay funciones disponibles' aunque están creadas?")
body("El sistema oculta automáticamente las funciones cuya hora ya pasó durante el día. Por ejemplo, si son las 8 PM, una función de las 7 PM ya no aparece. Esto es normal y esperado.")

h3("¿Un cliente dice que sus asientos 'ya están ocupados' pero nunca pagó?")
body("Cuando alguien selecciona sillas, quedan reservadas por 5 minutos. Si el timer expiró, se liberan automáticamente. Si el problema persiste, contactar al desarrollador — podría ser un hold que quedó atascado en Score.")

h2("5.2  Preguntas del panel de administración")

h3("¿No puedo entrar al /admin?")
body("Verificar que:")
bullet("Está iniciando sesión con el correo correcto (el autorizado como admin)")
bullet("Tiene conexión a internet estable")
bullet("Intente borrar cookies del navegador y volver a iniciar sesión")

h3("Las estadísticas del Dashboard muestran 0 en todo")
body("Esto puede pasar si Supabase (la base de datos) no está respondiendo. Recargar la página. Si persiste, contactar al desarrollador.")

h3("La sincronización con Score falla")
body("Para que Score funcione:")
bullet("El computador en el cine debe estar encendido y conectado a internet")
bullet("El programa ngrok debe estar corriendo (pantalla negra con texto)")
bullet("Si ngrok está cerrado, abrirlo con: C:\\ngrok\\ngrok.exe http 80")
nota("Si no sabe cómo revisar ngrok, llamar al desarrollador — es un proceso de un minuto.", "info")

h2("5.3  Preguntas de pagos y MercadoPago")

h3("¿Cuándo llega el dinero de las ventas a la cuenta del cine?")
body("Depende de la configuración de retiros en MercadoPago:")
bullet("Retiro automático diario: el dinero llega al día siguiente hábil")
bullet("Retiro manual: cuando el administrador lo solicite desde la app de MercadoPago")
nota("MercadoPago puede retener pagos mientras la cuenta está en proceso de verificación. Una vez verificada, los pagos fluyen normalmente.", "warn")

h3("¿Cómo ver las ventas en MercadoPago?")
bullet("Entrar a mercadopago.com.co con la cuenta del cine")
bullet("Ir a 'Tu negocio' → 'Actividad' para ver el historial completo de cobros")
bullet("Exportar reportes desde 'Tu negocio' → 'Reportes'")

h3("Un cliente reporta que pagó pero no tiene boleta y no aparece en reservas")
body("Pasos a seguir:")
bullet("Pedir al cliente el comprobante de pago de MercadoPago (número de referencia)")
bullet("Buscar esa referencia en el panel Admin → Reservas")
bullet("Si no aparece, el pago puede estar en estado 'Pendiente' — contactar al desarrollador con la referencia de pago para investigar")

separador()

# ── Pie de página ─────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Documento elaborado por el equipo de desarrollo · CINEWORLD · Cali, Colombia · Junio 2026")
r.font.size = Pt(9)
r.font.color.rgb = GRIS_CLARO
r.font.name = 'Calibri'
r.italic = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Para soporte técnico contactar al desarrollador: brianl280499@gmail.com")
r2.font.size = Pt(9)
r2.font.color.rgb = AZUL
r2.font.name = 'Calibri'

# ── Guardar ───────────────────────────────────────────────────────────────────
output_path = "/mnt/c/Proyectos/cineworld-web/CineWorld_Manual_Admin.docx"
doc.save(output_path)
print(f"✅ Documento guardado en: {output_path}")
